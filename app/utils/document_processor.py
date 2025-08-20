"""Document processor for extracting text from various file formats."""

import io
import re
from typing import Optional
import PyPDF2
from docx import Document as DocxDocument

from app.core.exceptions import ValidationError
from app.core.logging import get_logger

logger = get_logger(__name__)


class DocumentProcessor:
    """Utility class for processing and extracting text from documents."""
    
    async def extract_text(self, file_content: bytes, content_type: str) -> str:
        """
        Extract text from file content based on content type.
        
        Args:
            file_content: File content as bytes
            content_type: MIME content type
            
        Returns:
            Extracted text content
            
        Raises:
            ValidationError: If file processing fails
        """
        try:
            if content_type == "application/pdf" or content_type.endswith("/pdf"):
                return await self._extract_pdf_text(file_content)
            elif content_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword"
            ]:
                return await self._extract_docx_text(file_content)
            elif content_type.startswith("text/"):
                return await self._extract_plain_text(file_content)
            else:
                # Try to extract as plain text for unknown types
                return await self._extract_plain_text(file_content)
                
        except Exception as e:
            logger.error("Document processing failed", content_type=content_type, error=str(e))
            raise ValidationError(f"Failed to process document: {str(e)}")
    
    async def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_content.append(page.extract_text())
            
            # Join pages and clean up text
            full_text = "\n".join(text_content)
            return self._clean_text(full_text)
            
        except Exception as e:
            raise ValidationError(f"Failed to extract PDF text: {str(e)}")
    
    async def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            docx_file = io.BytesIO(file_content)
            document = DocxDocument(docx_file)
            
            text_content = []
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in document.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            full_text = "\n".join(text_content)
            return self._clean_text(full_text)
            
        except Exception as e:
            raise ValidationError(f"Failed to extract DOCX text: {str(e)}")
    
    async def _extract_plain_text(self, file_content: bytes) -> str:
        """Extract text from plain text file."""
        try:
            # Try different encodings
            encodings = ["utf-8", "utf-16", "latin1", "cp1252"]
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    return self._clean_text(text)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            text = file_content.decode("utf-8", errors="replace")
            return self._clean_text(text)
            
        except Exception as e:
            raise ValidationError(f"Failed to extract text: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove control characters except newlines and tabs
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', text)
        
        # Normalize line breaks
        text = re.sub(r'\r\n|\r', '\n', text)
        
        # Remove excessive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def get_text_preview(self, text: str, max_length: int = 500) -> str:
        """
        Get a preview of the text content.
        
        Args:
            text: Full text content
            max_length: Maximum length of preview
            
        Returns:
            Text preview
        """
        if len(text) <= max_length:
            return text
        
        # Try to break at sentence boundaries
        preview = text[:max_length]
        
        # Find the last sentence ending
        last_sentence = max(
            preview.rfind('.'),
            preview.rfind('!'),
            preview.rfind('?')
        )
        
        if last_sentence > max_length * 0.7:  # If we found a good break point
            return preview[:last_sentence + 1]
        else:
            # Otherwise, break at word boundary
            last_space = preview.rfind(' ')
            if last_space > 0:
                return preview[:last_space] + "..."
            else:
                return preview + "..."
    
    def count_words(self, text: str) -> int:
        """
        Count words in text.
        
        Args:
            text: Text content
            
        Returns:
            Word count
        """
        if not text:
            return 0
        
        # Simple word count based on whitespace
        words = text.split()
        return len(words)
    
    def estimate_reading_time(self, text: str, words_per_minute: int = 200) -> int:
        """
        Estimate reading time in minutes.
        
        Args:
            text: Text content
            words_per_minute: Average reading speed
            
        Returns:
            Estimated reading time in minutes
        """
        word_count = self.count_words(text)
        return max(1, round(word_count / words_per_minute))
    
    def extract_metadata(self, text: str) -> dict:
        """
        Extract basic metadata from text.
        
        Args:
            text: Text content
            
        Returns:
            Dictionary with metadata
        """
        return {
            "word_count": self.count_words(text),
            "character_count": len(text),
            "line_count": text.count('\n') + 1 if text else 0,
            "estimated_reading_time": self.estimate_reading_time(text),
            "preview": self.get_text_preview(text, 200)
        }